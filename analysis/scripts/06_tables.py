#!/usr/bin/env python
"""Stage 6 — every manuscript table, written as CSV then rendered to DOCX.

Tables 2, 3 and 4 are produced by stages 2, 3 and 4. This stage builds Table 1, assembles the
supplementary tables, and renders the whole set into a single DOCX with grid borders and footnotes.

Rewritten 2026-08-03 after the table audit. Six things were wrong and are fixed here:
  * one '%.4g' formatter served every column, so decimals varied inside a column and 1.000 printed
    as a bare '1'. Formats are now declared per column, per table.
  * Table 4 printed 145.3 where the legend and the Results quote 145.3104, so the published
    function could not be reproduced from the published table.
  * the legends were hard-coded here AND written in sources/09_Table_legends.md, and had drifted.
    That file is now the only source; this stage parses it.
  * Table 2 shipped 34 columns for two rows, four of them design constants and two of them exact
    duplicates of another column. It is now rendered in long form, in three blocks.
  * the tolerance-class table was written as 'SuppTable_S1', a number the legend file reserves for
    the author-supplied genotype list. It is now S9, and S1 is the genotype list.
  * SuppTable_S5 shipped as five bare columns with its provenance recorded nowhere in the file, and
    with a genotype key that would not join to S4 or S9.

Outputs: results/tables/Table1_traits.csv, SuppTable_*.csv, results/tables/Tables.docx
"""
import math
import os
import re
import numpy as np
import pandas as pd

import figstyle as S

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
DERIVED = os.path.join(ROOT, "analysis", "data", "derived")
TABLES = os.path.join(ROOT, "analysis", "results", "tables")
LEGENDS = os.path.join(ROOT, "manuscript", "sources", "09_Table_legends.md")
os.makedirs(TABLES, exist_ok=True)

TRAIT_TABLE = [
    ("CCI_TL", "Chlorophyll concentration index, top leaf", "µmol m⁻²",
     "Apogee MC-100", "youngest fully expanded leaf, 10 cm below the tip", "0, 30, 60 DAT"),
    ("CCI_BL", "Chlorophyll concentration index, bottom leaf", "µmol m⁻²",
     "Apogee MC-100", "lowest green leaf, 10 cm below the tip", "0, 30, 60 DAT"),
    ("rCCI", "Ratio of bottom-leaf to top-leaf chlorophyll concentration index",
     "dimensionless", "derived", "CCI_BL / CCI_TL", "30, 60 DAT"),
    ("QY_TL", "Maximum quantum yield of photosystem II, top leaf", "dimensionless (Fv/Fm)",
     "FluorPen FP-100, dark-adapted", "youngest fully expanded leaf", "0, 30, 60 DAT"),
    ("QY_BL", "Maximum quantum yield of photosystem II, bottom leaf", "dimensionless (Fv/Fm)",
     "FluorPen FP-100, dark-adapted", "lowest green leaf", "0, 30, 60 DAT"),
    ("rQY", "Ratio of bottom-leaf to top-leaf quantum yield", "dimensionless",
     "derived", "QY_BL / QY_TL", "30, 60 DAT"),
    ("Photo", "Net photosynthetic rate", "µmol CO₂ m⁻² s⁻¹",
     "LI-6400XT",
     "1000 µmol m⁻² s⁻¹ PAR, 30 °C, 60% RH, "
     "400 µmol mol⁻¹ CO₂", "0, 30, 60 DAT"),
    ("Cond", "Stomatal conductance", "mol H₂O m⁻² s⁻¹", "LI-6400XT",
     "as above", "0, 30, 60 DAT"),
    ("Ci", "Intercellular CO2 concentration", "µmol CO₂ mol⁻¹", "LI-6400XT",
     "as above", "0, 30, 60 DAT"),
    ("Tr", "Transpiration rate", "mmol H₂O m⁻² s⁻¹", "LI-6400XT",
     "as above", "0, 30, 60 DAT"),
]


# ---------------------------------------------------------------- formatting
def fmt_p(v):
    """Exact P at two significant figures, in the notation the manuscript body uses."""
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return ""
    try:
        v = float(v)
    except (TypeError, ValueError):
        return str(v)
    if pd.isna(v):
        return ""
    # Below 0.01 a three-decimal rendering loses the leading digit (0.001172 -> "0.001") and stops
    # matching the "1.2 x 10^-3" the Results quotes for the same quantity.
    if v >= 0.01:
        return "%.3f" % v
    if v <= 0:
        return "0"
    e = int(math.floor(math.log10(v)))
    m = v / 10 ** e
    sup = str(e).replace("-", "⁻")
    for a, b in zip("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹"):
        sup = sup.replace(a, b)
    return "%.1f × 10%s" % (m, sup)


# One fixed decimal count per column, so a column never mixes two and six significant figures.
FMT = {
    "Table3_performance.csv": {
        "n_traits": "%d", "n": "%d", "correct": "%d",
        "accuracy": "%.3f", "acc_lo": "%.3f", "acc_hi": "%.3f",
        "sensitivity": "%.3f", "specificity": "%.3f", "ppv": "%.3f", "npv": "%.3f",
        "TP": "%d", "TN": "%d", "FP": "%d", "FN": "%d",
        "genotype_folds_perfect": "%d", "fold_lo": "%.3f", "fold_hi": "%.3f"},
    "Table4_discriminant.csv": {
        "coefficient": "%.4f", "standardised": "%.3f", "boot_lo": "%.3f", "boot_hi": "%.3f"},
    "Table2_anova.csv": {
        "mean": "%.3f", "sd": "%.4f", "cv_pct": "%.1f",
        "df": "%d", "MS": "%.5f", "F": "%.2f", "P": fmt_p,
        "Estimate": "%.6f", "% of σ²P": "%.1f"},
}

# Raw script identifiers are not table headers.
HEADERS = {
    "trait": "Trait", "model": "Rule", "term": "Term", "n_traits": "Traits", "n": "n",
    "correct": "Correct (of 90)", "accuracy": "Accuracy",
    "acc_lo": "95% CI lower", "acc_hi": "95% CI upper",
    "sensitivity": "Sensitivity", "specificity": "Specificity", "ppv": "PPV", "npv": "NPV",
    "genotype_folds_perfect": "Perfect folds (of 15)",
    "fold_lo": "Folds 95% CI lower", "fold_hi": "Folds 95% CI upper",
    "coefficient": "Coefficient (units of D)", "standardised": "Standardised coefficient",
    "boot_lo": "Bootstrap 2.5th", "boot_hi": "Bootstrap 97.5th",
    "mean": "Mean", "sd": "SD", "cv_pct": "CV (%)",
}

# Symbols that are italic in running text and must be italic in a header too.
ITALIC = {"n", "F", "P", "D", "d"}

FOOTNOTES = {
    "Table1_traits.csv":
        "DAT, days after transplanting. The chlorophyll concentration index is in µmol m⁻², "
        "the unit the instrument reports; rCCI and rQY are dimensionless because the units "
        "cancel in the ratio; Fv/Fm is "
        "dimensionless and bounded at 1.",
    "Table2_anova.csv":
        "df, degrees of freedom; MS, mean square; σ²P, phenotypic variance; "
        "H², broad-sense heritability. n = 90 (15 genotypes × 2 nitrogen regimes "
        "× 3 replicate plants). Both traits are dimensionless, so MS carries trait units "
        "squared. The error term is the residual mean square on 60 degrees of freedom. Exact P "
        "values are given and no significance markers are used. The method-of-moments estimate "
        "of σ²G for QY_BL_30 was −1.095 × 10⁻⁵ before truncation "
        "and is reported as zero. The full 34-column computation, including the untruncated "
        "components, is in analysis/results/tables/Table2_anova.csv.",
    "Table3_performance.csv":
        "PPV, positive predictive value; NPV, negative predictive value; TP, TN, FP, FN, counts "
        "of true and false positives and negatives. n = 90 pooled held-out predictions over 15 "
        "folds. Intervals are exact (Clopper-Pearson). The majority-class baseline is 45 of 90.",
    "Table4_discriminant.csv":
        "Em dash, not applicable: the constant has no within-class standardised form and no "
        "bootstrap interval. Intervals are 95% percentile bootstrap over 5000 resamples of "
        "plants. Coefficients are printed at the precision needed to reproduce the boundary.",
}


def _fmt_cell(name, col, v):
    if not isinstance(v, str) and pd.isna(v):
        return "—"
    f = FMT.get(name, {}).get(col)
    if f is None:
        return ("%.4g" % v) if isinstance(v, (int, float, np.floating)) else str(v)
    if callable(f):
        return f(v)
    if f == "%d":
        return f % int(round(float(v)))
    return f % float(v)


# The genotype naming table lives in figstyle, so the tables and the figures cannot drift apart
# on a spelling. This file used to carry its own copy of it.
gkey = S.key


# ---------------------------------------------------------------- main tables
def table1():
    df = pd.DataFrame(TRAIT_TABLE, columns=[
        "Abbreviation", "Trait", "Unit", "Instrument", "Measurement position or formula", "Timing"])
    df.to_csv(os.path.join(TABLES, "Table1_traits.csv"), index=False)
    print("  Table1_traits.csv        %d rows" % len(df))
    return df


def table2_blocks():
    """Table 2 as three long-form blocks. The wide CSV keeps every computed column; the printed
    table gives the reader the descriptives, then df, MS, F and P per term, then the components."""
    p = os.path.join(TABLES, "Table2_anova.csv")
    if not os.path.exists(p):
        return []
    w = pd.read_csv(p)
    desc = w[["trait", "mean", "sd", "cv_pct"]].copy()

    terms = [("Genotype", "df_G", "MS_G", "F_G", "p_G"),
             ("Nitrogen", "df_T", "MS_T", "F_T", "p_T"),
             ("Genotype × nitrogen", "df_GxT", "MS_GxT", "F_GxT", "p_GxT"),
             ("Residual", "df_e", "MS_e", None, None)]
    rows = []
    for _, r in w.iterrows():
        for lab, dfc, msc, fc, pc in terms:
            rows.append({"trait": r["trait"], "Source of variation": lab,
                         "df": r[dfc], "MS": r[msc],
                         "F": np.nan if fc is None else r[fc],
                         "P": np.nan if pc is None else r[pc]})
    anova = pd.DataFrame(rows)

    comps = [("σ²G (genotype)", "var_G", "pct_var_G"),
             ("σ²T (nitrogen)", "var_T", "pct_var_T"),
             ("σ²G×T (interaction)", "var_GxT", "pct_var_GxT"),
             ("σ²e (residual)", "var_e", "pct_var_e"),
             ("σ²P (total)", "var_P", None),
             ("H² (broad sense)", "H2_broad", None)]
    rows = []
    for _, r in w.iterrows():
        for lab, vc, pc in comps:
            rows.append({"trait": r["trait"], "Variance component": lab,
                         "Estimate": r[vc],
                         "% of σ²P": np.nan if pc is None else r[pc]})
    var = pd.DataFrame(rows)
    return [(desc, "descriptive statistics"),
            (anova, "analysis of variance"),
            (var, "variance components")]


# ------------------------------------------------------------ supplementary
def supp_tables():
    made = []

    # --- S2 soil ----------------------------------------------------------
    soil = pd.DataFrame([
        # Named as Methods names them, with the original reference in brackets. Toth and Prince
        # (1949) IS the flame-photometric determination and Datta et al. (1962) IS the rapid
        # colorimetric one, so three of the four apparent contradictions with Methods were two
        # names for one method. Available nitrogen is the one that genuinely differs.
        ("Available N", "Alkaline permanganate (Subbiah and Asija 1956)", "kg ha-1", 220.62, 155.17),
        ("Available P2O5", "Olsen et al. (1982)", "kg ha-1", 96.06, 67.19),
        ("Available K2O", "Flame photometry (Toth and Prince 1949)", "kg ha-1", 162.42, 113.15),
        # The authors name Datta et al. (1962), not Walkley and Black. Restored 2026-08-03 after an
        # automated rewrite of this file reverted it.
        ("Organic carbon", "Rapid colorimetry (Datta et al. 1962)", "%", 0.69, 0.58),
        ("pH", "1:2.5 soil:water", "", 7.29, 7.20),
        ("Electrical conductivity", "1:2.5 soil:water", "dS m-1", 0.24, 0.22),
    ], columns=["Soil property", "Method", "Unit", "Control", "Low nitrogen"])
    p = os.path.join(TABLES, "SuppTable_S2_soil_properties.csv")
    # Without float_format the pH row printed 7.29 against 7.2 in the same column.
    soil.to_csv(p, index=False, float_format="%.2f"); made.append(p)

    # --- S3 trait dictionary ---------------------------------------------
    src = os.path.join(DERIVED, "trait_dictionary.tsv")
    if os.path.exists(src):
        td = pd.read_csv(src, sep="\t")
        # Table 1 says "dimensionless" and the dictionary said "no units" for the same concept.
        td["unit"] = td["unit"].replace({"no units": "dimensionless"})
        p = os.path.join(TABLES, "SuppTable_S3_trait_dictionary.csv")
        td.to_csv(p, index=False); made.append(p)

    gm = pd.read_csv(os.path.join(DERIVED, "genotype_means.tsv"), sep="\t")
    plants = pd.read_csv(os.path.join(DERIVED, "plants_30DAT.tsv"), sep="\t")
    code = (gm[["genotype_code", "genotype_name"]].drop_duplicates()
            .assign(k=lambda d: d["genotype_name"].map(gkey))
            .drop_duplicates("k").set_index("k")["genotype_code"])

    # --- S4 genotype means ------------------------------------------------
    focal = gm[(gm["timepoint"] == "30 DAT") &
               (gm["trait"].isin(["CCI_TL", "CCI_BL", "rCCI", "QY_TL", "QY_BL", "rQY"]))]
    wide = focal.pivot_table(index=["genotype_code", "genotype_name", "treatment"],
                             columns="trait", values=["value", "se"])
    wide.columns = ["%s (%s)" % (t, "mean" if a == "value" else "s.e.")
                    for a, t in wide.columns]
    wide = wide.reset_index()
    # The published means carry two decimal places. Writing them at three fabricated a digit, and
    # the se column that genotype_means.tsv carries was being dropped entirely.
    for c in wide.columns:
        if wide[c].dtype.kind == "f":
            wide[c] = wide[c].map(lambda v: v if pd.isna(v) else round(float(v), 2))
    # The Results quote means recomputed from the transcribed plant matrix of Table S5, which the
    # published two-decimal means do not reproduce exactly. Both are shipped, side by side.
    rec = (plants.assign(k=plants["genotype"].map(gkey))
           .groupby(["k", "treatment"])[["rCCI_30", "QY_BL_30"]].mean().round(3).reset_index())
    wide["k"] = wide["genotype_name"].map(gkey)
    wide = wide.merge(rec, on=["k", "treatment"], how="left").drop(columns="k")
    wide = wide.rename(columns={"rCCI_30": "rCCI_30 (recomputed from Table S5)",
                                "QY_BL_30": "QY_BL_30 (recomputed from Table S5)"})
    p = os.path.join(TABLES, "SuppTable_S4_genotype_means_30DAT.csv")
    wide.to_csv(p, index=False); made.append(p)

    # --- S5 plant level ---------------------------------------------------
    pl = plants.copy()
    pl.insert(0, "genotype_code", pl["genotype"].map(gkey).map(code))
    p = os.path.join(TABLES, "SuppTable_S5_plant_level_30DAT.csv")
    # A reader who opens this file alone sees 90 rows of two-decimal numbers and nothing that says
    # where they came from. Read it back with pd.read_csv(path, comment='#').
    with open(p, "w", encoding="utf-8", newline="") as f:
        f.write("# Table S5. Plant-level rCCI_30 and QY_BL_30, 30 days after transplanting, "
                "n = 90.\n")
        f.write("# PROVENANCE: recovered by cell-by-cell transcription from the rendered heat-map "
                "figures of the source package (ref 29); the source matrix has not been "
                "released.\n")
        f.write("# RESOLUTION: every value carries two decimal places; each number stands for an "
                "interval of width 0.01.\n")
        f.write("# Sensitivity of every downstream result to that rounding is in Table S6.\n")
        pl.to_csv(f, index=False)
    made.append(p)

    # --- S6, S7, S8: cited in Methods and Results, previously shipped only under raw names ---
    for n, stem in ((6, "robustness_checks"), (7, "trait_pair_decision"),
                    (8, "within_treatment_genotype")):
        raw = os.path.join(TABLES, stem + ".csv")
        if not os.path.exists(raw):
            continue
        d = pd.read_csv(raw)
        for c in d.columns:
            if c.lower() in ("p", "corr_p", "p_genotype") or c.lower().startswith("p_value"):
                d[c] = d[c].map(fmt_p)
        p = os.path.join(TABLES, "SuppTable_S%d_%s.csv" % (n, stem))
        d.to_csv(p, index=False); made.append(p)

    # --- S1 the fifteen genotypes ----------------------------------------
    # Transcribed from the authors' own "Supplementary tables.docx", table 1 (Plant ID, Main
    # Population, Genotype Name, Origin), which the register verified cell by cell against
    # thesis Table 3.1 p44. It is the one table in their package adopted as given. Only the
    # spelling of the names is normalised, to the single naming table in figstyle, so that this
    # table joins to S4, S5 and S9; the authors' own spelling is kept alongside it.
    author_list = [
        ("GT001", "Japonica", "Moroberakan", "Guinea"),
        ("GT002", "Indica", "IR83388-B-B108-3", "India"),
        ("GT003", "Indica", "IR77298-14-1-2-10", "India"),
        ("GT004", "Basmati", "PUSA1121", "India"),
        ("GT005", "Basmati", "BAM8315", "India"),
        ("GT006", "Basmati", "BAM812", "India"),
        ("GT007", "Aus", "Malchi", "India"),
        ("GT008", "Basmati", "BAM 3690", "India"),
        ("GT009", "Basmati", "BAM 4138", "India"),
        ("GT010", "Aus", "BAM 4521", "India"),
        ("GT011", "Indica", "Black gora", "India"),
        ("GT012", "Japonica", "Suweon", "Republic of Korea"),
        ("GT013", "Indica", "Kunjukunju", "India"),
        ("GT014", "Indica", "Cauvery", "India"),
        ("GT015", "Indica", "RPW9-4(SS1)", "India"),
    ]
    g1 = pd.DataFrame(author_list, columns=["Plant ID", "Subspecies group",
                                            "Name as published by the authors", "Origin"])
    g1.insert(2, "Genotype", g1["Name as published by the authors"].map(S.disp))
    g1.insert(0, "Genotype code", g1["Name as published by the authors"].map(gkey).map(code))
    g1 = g1[["Genotype code", "Plant ID", "Genotype", "Name as published by the authors",
             "Subspecies group", "Origin"]]
    r = os.path.join(TABLES, "SuppTable_S1_genotype_list.csv")
    g1.to_csv(r, index=False); made.append(r)

    # --- S9 tolerance classes --------------------------------------------
    gl = pd.read_csv(os.path.join(DERIVED, "genotype_tolerance.csv"))
    gl["genotype_code"] = gl["genotype_name"].map(gkey).map(code)
    # "thesis Table 4.26 p132" is a pointer into an unreleased internal document; a reader of the
    # supplement cannot follow it.
    gl["source"] = "Published for this experiment [29]"
    gl = gl.rename(columns={"genotype_code": "Genotype code", "genotype_name": "Genotype",
                            "tolerance_class": "Tolerance class",
                            "cluster_Nstress_AL": "Cluster (average linkage, low N)",
                            "source": "Source"})
    gl = gl[["Genotype code", "Genotype", "Cluster (average linkage, low N)",
             "Tolerance class", "Source"]]
    # This file used to be written as "SuppTable_S1"; that number belongs to the genotype list.
    q = os.path.join(TABLES, "SuppTable_S9_tolerance_classes.csv")
    gl.to_csv(q, index=False); made.append(q)

    for m in made:
        print("  %-46s" % os.path.basename(m))
    return made


# ---------------------------------------------------------------- rendering
def read_legends():
    """09_Table_legends.md is the master. Parsing it means the DOCX and the manuscript cannot
    carry two different legends for the same table, which they did: one said '3 replicate pots'
    and the other '3 replicate plants' for the same n = 90."""
    out = {}
    if not os.path.exists(LEGENDS):
        return out
    txt = open(LEGENDS, encoding="utf-8").read()
    txt = re.sub(r"<!--.*?-->", "", txt, flags=re.S)
    for line in txt.splitlines():
        m = re.match(r"^\*\*Table (\d)\.", line.strip())
        if m:
            clean = re.sub(r"\{[^}]*\}", "", line.strip())
            clean = clean.replace("**", "").replace("*", "")
            out["Table%s" % m.group(1)] = re.sub(r"\s+", " ", clean).strip()
    return out


def table3_long():
    """Table 3 with the rules as columns and the metrics as rows.

    As written it had eighteen columns for three rules. On A4 portrait that gives about 9 mm a
    column, so every header broke over three or four lines ("Cor / rect / (of / 90)") and values
    split mid-number. Three rules transpose to three columns and read at full size.
    """
    d = pd.read_csv(os.path.join(TABLES, "Table3_performance.csv"))
    rows = [
        ("Traits in the rule", "n_traits", "%d"),
        ("Plants classified correctly (of 90)", "correct", "%d"),
        ("Accuracy", "accuracy", "%.3f"),
        ("Accuracy, exact 95% CI", ("acc_lo", "acc_hi"), None),
        ("Sensitivity", "sensitivity", "%.3f"),
        ("Specificity", "specificity", "%.3f"),
        ("Positive predictive value", "ppv", "%.3f"),
        ("Negative predictive value", "npv", "%.3f"),
        ("True positives", "TP", "%d"),
        ("True negatives", "TN", "%d"),
        ("False positives", "FP", "%d"),
        ("False negatives", "FN", "%d"),
        ("Genotype folds fully correct (of 15)", "genotype_folds_perfect", "%d"),
        ("Folds, exact 95% CI", ("fold_lo", "fold_hi"), None),
    ]
    out = []
    for label, col, fmt in rows:
        rec = {"Measure": label}
        for _, r in d.iterrows():
            name = MODEL_LABEL.get(r["model"], r["model"])
            if isinstance(col, tuple):
                rec[name] = "%.3f to %.3f" % (r[col[0]], r[col[1]])
            else:
                rec[name] = fmt % r[col]
        out.append(rec)
    return pd.DataFrame(out)


MODEL_LABEL = {"rCCI_30": "rCCI_30 alone", "QY_BL_30": "QY_BL_30 alone", "both": "Both traits"}


ORDER_MAIN = {1: "Table1_traits.csv", 2: "Table2_anova.csv",
              3: "Table3_performance.csv", 4: "Table4_discriminant.csv"}


def append_main_table(doc, n):
    """Append main table `n` to an existing Document, WITHOUT its legend.

    build_manuscript.py calls this so each table sits directly under the legend that already
    stands in 09_Table_legends.md, which is where BMC wants tables — in the manuscript file, not
    as loose attachments. The four tables were being written to analysis/results/tables/Tables.docx
    and never copied into the submission bundle, so the manuscript cited Table 1 to Table 4 and
    shipped none of them. Found 2026-08-03.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    name = ORDER_MAIN.get(n)
    path = os.path.join(TABLES, name) if name else None
    if not path or not os.path.exists(path):
        return False

    def write_table(nm, df):
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = "Table Grid"
        for i, c in enumerate(df.columns):
            head = HEADERS.get(c, c)
            cell = t.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(head)
            run.bold = True
            run.italic = head in ITALIC
        for _, r in df.iterrows():
            cells = t.add_row().cells
            for i, c in enumerate(df.columns):
                cells[i].text = _fmt_cell(nm, c, r[c])
        # The manuscript's Normal style is 12 pt double-spaced, which a table inherits and which
        # would push a ten-column table over several pages of air. Tables are set single-spaced.
        for row in t.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    par.paragraph_format.space_after = Pt(0)
                    par.paragraph_format.first_line_indent = Pt(0)
                    for run in par.runs:
                        run.font.size = Pt(8)
                        run.font.name = "Times New Roman"

    if name == "Table3_performance.csv":
        write_table(name, table3_long())
    elif name == "Table2_anova.csv":
        for d, lab in table2_blocks():
            lp = doc.add_paragraph()
            lp.paragraph_format.first_line_indent = Pt(0)
            lp.add_run(lab).italic = True
            write_table(name, d)
            doc.add_paragraph("")
    else:
        write_table(name, pd.read_csv(path))
    fn = FOOTNOTES.get(name)
    if fn:
        fp = doc.add_paragraph()
        fp.paragraph_format.first_line_indent = Pt(0)
        fp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fp.add_run(fn).font.size = Pt(8)
    doc.add_paragraph("")
    return True


def render_docx():
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("  (python-docx unavailable — CSV only)")
        return
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    leg = read_legends()

    def write_table(name, df):
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = "Table Grid"
        for i, c in enumerate(df.columns):
            head = HEADERS.get(c, c)
            cell = t.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(head)
            run.bold = True
            run.italic = head in ITALIC
        for _, r in df.iterrows():
            cells = t.add_row().cells
            for i, c in enumerate(df.columns):
                cells[i].text = _fmt_cell(name, c, r[c])

    order = ["Table1_traits.csv", "Table2_anova.csv", "Table3_performance.csv",
             "Table4_discriminant.csv"]
    for n, name in enumerate(order, 1):
        p = os.path.join(TABLES, name)
        if not os.path.exists(p):
            continue
        doc.add_paragraph(leg.get("Table%d" % n,
                                  "Table %d. {AUTHOR: legend missing from "
                                  "09_Table_legends.md}" % n))
        if name == "Table3_performance.csv":
            write_table(name, table3_long())
        elif name == "Table2_anova.csv":
            for d, lab in table2_blocks():
                doc.add_paragraph(lab).runs[0].italic = True
                write_table(name, d)
                doc.add_paragraph("")
        else:
            write_table(name, pd.read_csv(p))
        fn = FOOTNOTES.get(name)
        if fn:
            doc.add_paragraph().add_run(fn).font.size = Pt(8)
        doc.add_paragraph("")

    out = os.path.join(TABLES, "Tables.docx")
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(out)
    print("  Tables.docx")


def main():
    print("tables ->", TABLES)
    table1()
    supp_tables()
    render_docx()


if __name__ == "__main__":
    main()
